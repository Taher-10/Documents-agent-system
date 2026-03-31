"""conftest.py — add src/ to sys.path so tests import document_parser directly."""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

import fitz
import pytest
from docx import Document

@pytest.fixture()
def pdf_path(tmp_path: Path) -> Path:
    """Two-page PDF: page 1 is text-rich (>50 chars), page 2 is near-empty."""
    doc = fitz.open()

    p1 = doc.new_page()
    p1.insert_text(
        (72, 72),
        "Objet\nCe document décrit la procédure de gestion des gabarits.\n" * 4,
    )

    p2 = doc.new_page()
    p2.insert_text((72, 72), "ok")  # <50 chars → triggers fitz fallback

    out = tmp_path / "test.pdf"
    out.write_bytes(doc.tobytes())
    doc.close()
    return out


@pytest.fixture()
def docx_path(tmp_path: Path) -> Path:
    """DOCX with a Heading 1, Heading 2, a normal paragraph, and a 2x2 table."""
    doc = Document()

    doc.add_heading("Objet", level=1)
    doc.add_heading("Domaine d'application", level=2)
    doc.add_paragraph("Ce document décrit la procédure de gestion des gabarits.")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Colonne A"
    table.cell(0, 1).text = "Colonne B"
    table.cell(1, 0).text = "Valeur 1"
    table.cell(1, 1).text = "Valeur 2"

    out = tmp_path / "test.docx"
    doc.save(str(out))
    return out
