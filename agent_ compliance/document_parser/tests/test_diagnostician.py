"""test_diagnostician.py — M1 unit tests for document_diagnostician."""

import pytest
import fitz  # PyMuPDF
from pathlib import Path
from document_parser.document_diagnostician import PageInfo, PageMap, classify_page_type, assign_quality_tier, inspect_document
from document_parser.parsed_document import UnsupportedFormatError

# ---------------------------------------------------------------------------
# Task 1 — Dataclasses
# ---------------------------------------------------------------------------

def test_page_info_construction():
    info = PageInfo(
        page_number=1,
        page_type="text",
        has_selectable_text=True,
        image_count=0,
        font_issue=False,
        text_sample="Procédure de gestion",
    )
    assert info.page_number == 1
    assert info.page_type == "text"
    assert info.has_selectable_text is True
    assert info.image_count == 0
    assert info.font_issue is False
    assert info.text_sample == "Procédure de gestion"


def test_page_info_image_page():
    info = PageInfo(
        page_number=3,
        page_type="image",
        has_selectable_text=False,
        image_count=4,
        font_issue=False,
        text_sample="",
    )
    assert info.page_type == "image"
    assert info.has_selectable_text is False
    assert info.image_count == 4


def test_page_map_construction():
    page = PageInfo(
        page_number=1,
        page_type="text",
        has_selectable_text=True,
        image_count=0,
        font_issue=False,
        text_sample="Hello",
    )
    pm = PageMap(
        total_pages=1,
        file_format="pdf",
        quality_tier="A",
        pages=[page],
        producer="Microsoft Word 2016",
    )
    assert pm.total_pages == 1
    assert pm.file_format == "pdf"
    assert pm.quality_tier == "A"
    assert len(pm.pages) == 1
    assert pm.producer == "Microsoft Word 2016"


def test_page_map_producer_none():
    page = PageInfo(page_number=1, page_type="text", has_selectable_text=True, image_count=0, font_issue=False, text_sample="x" * 150)
    pm = PageMap(total_pages=1, file_format="pdf", quality_tier="A", pages=[page], producer=None)
    assert pm.producer is None


# ---------------------------------------------------------------------------
# Task 2 — classify_page_type
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("text,image_count,font_issue,expected", [
    # text: stripped len > 100 AND no font issue
    ("A" * 150, 0, False, "text"),
    ("Procédure " * 15, 2, False, "text"),   # len > 100, font OK → text even with images
    # image: stripped len < 30
    ("", 4, False, "image"),
    ("   ", 0, False, "image"),             # whitespace only → stripped len 0
    ("Short.", 3, False, "image"),           # 6 chars < 30
    ("A" * 29, 0, False, "image"),          # exactly 29 < 30
    # hybrid: everything else
    ("A" * 30, 0, False, "hybrid"),         # exactly 30 — boundary, not < 30 and not > 100
    ("A" * 100, 0, False, "hybrid"),        # exactly 100 — not > 100
    ("A" * 150, 0, True, "hybrid"),         # len > 100 but font issue
    ("A" * 60, 1, False, "hybrid"),         # 30 <= len <= 100
    ("A" * 60, 0, True, "hybrid"),          # 30 <= len, font issue
])
def test_classify_page_type(text, image_count, font_issue, expected):
    assert classify_page_type(text, image_count, font_issue) == expected


# ---------------------------------------------------------------------------
# Task 3 — assign_quality_tier
# ---------------------------------------------------------------------------


def _make_page_info(page_type: str, font_issue: bool = False, page_number: int = 1) -> PageInfo:
    """Helper: build a PageInfo with the given page_type."""
    return PageInfo(
        page_number=page_number,
        page_type=page_type,  # type: ignore[arg-type]
        has_selectable_text=(page_type == "text"),
        image_count=0,
        font_issue=font_issue,
        text_sample="",
    )


def test_assign_tier_a_all_text_no_font_issues():
    pages = [_make_page_info("text") for _ in range(5)]
    assert assign_quality_tier(pages) == "A"


def test_assign_tier_b_has_font_issue():
    pages = [_make_page_info("text", font_issue=True)] + [_make_page_info("text") for _ in range(4)]
    assert assign_quality_tier(pages) == "B"


def test_assign_tier_b_has_hybrid_page():
    pages = [_make_page_info("hybrid")] + [_make_page_info("text") for _ in range(4)]
    assert assign_quality_tier(pages) == "B"


def test_assign_tier_c_majority_image():
    # 3 image out of 5 → 60% > 50%
    pages = [_make_page_info("image")] * 3 + [_make_page_info("text")] * 2
    assert assign_quality_tier(pages) == "C"


def test_assign_tier_b_exactly_half_image():
    # 2 image out of 4 → exactly 50%, not > 50% → B
    pages = [_make_page_info("image")] * 2 + [_make_page_info("text")] * 2
    assert assign_quality_tier(pages) == "B"


def test_assign_tier_c_all_image():
    pages = [_make_page_info("image") for _ in range(4)]
    assert assign_quality_tier(pages) == "C"


def test_assign_tier_b_single_image_many_text():
    # 1 image out of 10 → 10%, not > 50% → B (not all text)
    pages = [_make_page_info("image")] + [_make_page_info("text")] * 9
    assert assign_quality_tier(pages) == "B"


# ---------------------------------------------------------------------------
# Task 4 — inspect_document: DOCX branch + UnsupportedFormatError
# ---------------------------------------------------------------------------


@pytest.fixture
def minimal_docx(tmp_path: Path) -> Path:
    """Minimal DOCX with two paragraphs of text."""
    from docx import Document as DocxDocument  # local import — optional dep
    path = tmp_path / "test_doc.docx"
    doc = DocxDocument()
    doc.add_paragraph("Ceci est un document de procédure QHSE.")
    doc.add_paragraph("Il contient plusieurs paragraphes de texte de test.")
    doc.save(str(path))
    return path


def test_inspect_docx_returns_tier_a(minimal_docx: Path):
    page_map = inspect_document(minimal_docx)
    assert page_map.file_format == "docx"
    assert page_map.quality_tier == "A"
    assert page_map.total_pages == 1
    assert len(page_map.pages) == 1


def test_inspect_docx_page_is_text(minimal_docx: Path):
    page_map = inspect_document(minimal_docx)
    page = page_map.pages[0]
    assert page.page_type == "text"
    assert page.has_selectable_text is True
    assert page.font_issue is False
    assert page.page_number == 1


def test_inspect_docx_producer_is_none(minimal_docx: Path):
    page_map = inspect_document(minimal_docx)
    assert page_map.producer is None


def test_inspect_docx_text_sample_captured(minimal_docx: Path):
    page_map = inspect_document(minimal_docx)
    assert "procédure QHSE" in page_map.pages[0].text_sample


def test_inspect_unsupported_format_raises(tmp_path: Path):
    bad_file = tmp_path / "document.txt"
    bad_file.write_text("not a pdf")
    with pytest.raises(UnsupportedFormatError, match=r"\.txt"):
        inspect_document(bad_file)


# ---------------------------------------------------------------------------
# Task 5 — inspect_document: PDF branch
# ---------------------------------------------------------------------------


@pytest.fixture
def text_pdf(tmp_path: Path) -> Path:
    """Single-page clean-text PDF (Tier A)."""
    path = tmp_path / "clean_text.pdf"
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)  # A4
    # Insert > 100 chars of selectable text
    page.insert_text(
        (72, 100),
        "Procédure de Gestion des Gabarits Externes. "
        "Ce document décrit les étapes nécessaires pour la gestion "
        "des gabarits utilisés dans le processus de production.",
        fontsize=11,
    )
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def image_pdf(tmp_path: Path) -> Path:
    """Single-page image-only PDF (Tier C via >50% image pages)."""
    path = tmp_path / "image_only.pdf"
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    # Insert a raster image covering the page — no selectable text
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 595, 842))
    pix.clear_with(200)  # fill with light gray
    page.insert_image(fitz.Rect(0, 0, 595, 842), pixmap=pix)
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def mixed_pdf(tmp_path: Path) -> Path:
    """Two-page PDF: page 1 = text, page 2 = image (Tier B)."""
    path = tmp_path / "mixed.pdf"
    doc = fitz.open()
    # Page 1 — text
    p1 = doc.new_page(width=595, height=842)
    p1.insert_text(
        (72, 100),
        "Objet: Ce document décrit la procédure de gestion. " * 4,
        fontsize=11,
    )
    # Page 2 — image only
    p2 = doc.new_page(width=595, height=842)
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 595, 842))
    pix.clear_with(200)
    p2.insert_image(fitz.Rect(0, 0, 595, 842), pixmap=pix)
    doc.save(str(path))
    doc.close()
    return path


def test_inspect_text_pdf_tier_a(text_pdf: Path):
    page_map = inspect_document(text_pdf)
    assert page_map.file_format == "pdf"
    assert page_map.quality_tier == "A"
    assert page_map.total_pages == 1
    assert len(page_map.pages) == 1


def test_inspect_text_pdf_page_classified_as_text(text_pdf: Path):
    page_map = inspect_document(text_pdf)
    page = page_map.pages[0]
    assert page.page_type == "text"
    assert page.has_selectable_text is True
    assert page.page_number == 1


def test_inspect_text_pdf_sample_captured(text_pdf: Path):
    page_map = inspect_document(text_pdf)
    assert "Procédure" in page_map.pages[0].text_sample


def test_inspect_image_pdf_tier_c(image_pdf: Path):
    page_map = inspect_document(image_pdf)
    assert page_map.quality_tier == "C"
    assert page_map.pages[0].page_type == "image"
    assert page_map.pages[0].has_selectable_text is False


def test_inspect_mixed_pdf_tier_b(mixed_pdf: Path):
    page_map = inspect_document(mixed_pdf)
    assert page_map.quality_tier == "B"
    assert page_map.total_pages == 2
    page_types = {p.page_number: p.page_type for p in page_map.pages}
    assert page_types[1] == "text"
    assert page_types[2] == "image"


def test_inspect_pdf_image_count(image_pdf: Path):
    page_map = inspect_document(image_pdf)
    assert page_map.pages[0].image_count >= 1
