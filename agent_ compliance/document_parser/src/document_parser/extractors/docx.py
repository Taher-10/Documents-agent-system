"""
docx.py — M2-DOCX extractor: DOCX files (python-docx).

The entire document is treated as a single logical page (page 1).
Paragraphs and tables are traversed in XML body order to preserve
the interleaved structure of QHSE documents (metadata table at top,
procedure text in middle, record-form table at bottom).

Heading paragraphs are prefixed with ##H{level}## for the M4 segmenter.
Table cells are stored structured in RawPageText.tables AND as a flat
text dump in RawPageText.text (for M3/M4 visibility until LLM table
transformation is implemented).

Implementation notes:
- `Paragraph(child, doc)` and `Table(child, doc)` use python-docx internal
  constructors (not the public `doc.paragraphs`/`doc.tables` accessors). This
  is intentional: iterating `doc.element.body` directly is the only way to
  preserve the interleaved order of paragraphs and tables in XML document order.

Known limitations:
- Merged cells: python-docx repeats the same cell object for horizontally
  merged cells. The flat text dump and `tables` field will contain duplicate
  cell text for merged-cell tables (common in QHSE header tables).
- Heading detection uses `style.name.startswith("Heading")`, which works for
  English-locale Word files. French-locale installations name styles "Titre 1",
  "Titre 2", etc. — these will be treated as normal paragraphs without a
  ##H{n}## prefix.
"""

from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from document_parser.parsed_document import RawPageText


def extract_docx(path: Path) -> list[RawPageText]:
    """Extract text and tables from a DOCX file.

    The entire document is returned as a single RawPageText (page_number=1).
    Elements are processed in document body order (XML traversal) to preserve
    the interleaved sequence of paragraphs and tables.

    Args:
        path: Absolute path to the .docx file.

    Returns:
        A one-element list containing one RawPageText for the whole document.
    """
    doc = Document(str(path))

    text_parts: list[str] = []
    tables: list[list[list[str]]] = []

    for child in doc.element.body:
        tag = child.tag.split("}")[-1]  # strip XML namespace prefix

        if tag == "p":
            para = Paragraph(child, doc)
            text = para.text.strip()
            if not text:
                continue  # skip empty paragraphs (common in Word docs)

            style = para.style.name  # e.g. "Heading 1", "Heading 2", "Normal"
            if style.startswith("Heading"):
                try:
                    level = int(style.split()[-1])
                except ValueError:
                    level = 1
                text_parts.append(f"##H{level}## {text}")
            else:
                text_parts.append(text)

        elif tag == "tbl":
            table = Table(child, doc)
            rows = [
                [cell.text.strip() for cell in row.cells]
                for row in table.rows
            ]
            tables.append(rows)
            # Flat dump so M3/M4 can see table content in the text stream
            flat = "\n".join("  ".join(cell for cell in row) for row in rows)
            text_parts.append(flat)

    return [
        RawPageText(
            page_number=1,
            text="\n".join(text_parts),
            tables=tables,
            extraction_method="docx",
            confidence=1.0,
        )
    ]
